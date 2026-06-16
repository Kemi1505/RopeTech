import os
import tempfile
import pandas as pd
from docx import Document
from mailmerge import MailMerge
import io
import win32com.client as win32
import pythoncom

LOAD_MAPPING_FILE = "shackles/shackle-data.xlsx"

def load_load_mapping(mapping_file):
    df = pd.read_excel(mapping_file)
    df = df.fillna('').astype(str)
    mapping = {}
    for _, row in df.iterrows():
        try:
            load = float(row.iloc[0])
        except:
            continue
        mapping[load] = row.to_dict()
    return mapping

def get_diagram_path(equipment, type_):
    equipment = equipment.lower().strip()
    type_ = type_.lower().replace('&', 'and').strip()
    if equipment == "bow shackle" and type_ == "bolt and nut":
        return "shackles/Bow-BnN.png"
    elif equipment == "bow shackle" and type_ == "screw pin":
        return "shackles/Bow-SP.png"
    elif equipment == "dee shackle" and type_ == "bolt and nut":
        return "shackles/Dee-BnN.png"
    elif equipment == "dee shackle" and type_ == "screw pin":
        return "shackles/Dee-SP.png"
    else:
        return "shackles/Bow-BnN.png"

def generate_shackle_certificates(excel_path, template_path, sheet_name, output_filename):
    if not os.path.exists(LOAD_MAPPING_FILE):
        raise FileNotFoundError(f"Load mapping file not found: {LOAD_MAPPING_FILE}")
    load_map = load_load_mapping(LOAD_MAPPING_FILE)

    df = pd.read_excel(excel_path, sheet_name=sheet_name)
    df = df.fillna('').astype(str)
    records = df.to_dict(orient='records')
    if not records:
        raise ValueError("No data rows found in shackle sheet.")

    print(f"📄 Generating {len(records)} certificates...")  # console debug

    temp_dir = tempfile.mkdtemp()
    temp_files = []

    for idx, row in enumerate(records):
        equipment = row.get('Equipment', 'bow shackle') or 'bow shackle'
        type_ = row.get('Type', 'bolt and nut') or 'bolt and nut'
        diagram_path = get_diagram_path(equipment, type_)
        if not os.path.exists(diagram_path):
            diagram_path = "shackles/Bow-BnN.png"

        pin_check = "N/A" if type_.lower() == "screw pin" else "ok"
        row['pinCheck'] = pin_check

        load_str = row.get('Load_Tonnes', '0') or '0'
        try:
            load = float(load_str)
        except ValueError:
            load = 0.0
        mapping_vals = load_map.get(load, load_map.get(list(load_map.keys())[0], {}))

        merge_data = {**row, **mapping_vals}

        with MailMerge(template_path) as mail_doc:
            mail_doc.merge(**merge_data)
            stream = io.BytesIO()
            mail_doc.write(stream)
            stream.seek(0)

        doc = Document(stream)
        for paragraph in doc.paragraphs:
            if "{Diagram}" in paragraph.text:
                paragraph.clear()
                run = paragraph.add_run()
                run.add_picture(diagram_path)
                break

        temp_path = os.path.join(temp_dir, f"cert_{idx:03d}.docx")
        doc.save(temp_path)
        temp_files.append(temp_path)
        print(f"   ✅ Saved cert #{idx+1}")  # console debug

    # --- Merge using Word COM ---
    pythoncom.CoInitialize()
    word = None
    try:
        word = win32.gencache.EnsureDispatch('Word.Application')
        word.Visible = False
        
        merged_doc = word.Documents.Add()
        selection = word.Selection
        
        for i, file_path in enumerate(temp_files):
            # Move cursor to the very end of the document
            selection.EndKey(Unit=6)   # wdStory
            # Insert the certificate file
            selection.InsertFile(file_path)
            # Add a page break after every certificate except the last
            if i < len(temp_files) - 1:
                selection.EndKey(Unit=6)   # move to end again
                selection.InsertBreak(Type=7)  # wdPageBreak

        output_path = os.path.join(temp_dir, output_filename)
        merged_doc.SaveAs(output_path)
        merged_doc.Close()
        print(f"✅ Merged {len(temp_files)} certificates into final document.")
    finally:
        if word is not None:
            word.Quit()
        pythoncom.CoUninitialize()

    return output_path
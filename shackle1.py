import os
import tempfile
import pandas as pd
from docx import Document
from docx.shared import Inches
import shutil

# Path to load mapping Excel (adjust as needed)
LOAD_MAPPING_FILE = "shackles/shackle-data.xlsx"

def load_load_mapping(mapping_file):
    """Read mapping Excel and return a dict keyed by load (float)."""
    df = pd.read_excel(mapping_file)
    df = df.fillna('').astype(str)
    mapping = {}
    for _, row in df.iterrows():
        try:
            load = float(row.iloc[0])  # assume first column is load value
        except:
            continue
        mapping[load] = row.to_dict()
    return mapping

def get_diagram_path(equipment, type_):
    """Return path to the correct diagram image based on equipment and type."""
    equipment = equipment.lower().strip()
    type_ = type_.lower().strip()
    
    # Define filenames (place these images in templates/ folder)
    if equipment == "bow shackle" and (type_ == "bolt and nut" or "bolt & nut"):
        return "shackles/Bow-BnN.png"
    elif equipment == "bow shackle" and type_ == "screw pin":
        return "shackles/Bow-SP.png"
    elif equipment == "dee shackle" and (type_ == "bolt and nut" or "bolt & nut"):
        return "shackles/Dee-BnN.png"
    elif equipment == "dee shackle" and type_ == "screw pin":
        return "shackles/Dee-SP.png"
    else:
        # fallback
        return "shackles/Bow-BnN.png"

def generate_shackle_certificates(excel_path, template_path, sheet_name, output_filename):
    """
    Generate certificates for shackles.
    
    Expected columns in data Excel (sheet_name):
      - Equipment: "bow shackle" or "dee shackle"
      - Type: "bolt and nut" or "screw pin"
      - Load_Tonnes: numeric value to look up in mapping
      - any other fields to merge (e.g., IdNo, VesselName, Date...)
    
    The template must contain placeholders:
      - «pinCheck» -> replaced with "N/A" or "ok"
      - «SWL», «Diameter», «Length» etc. (based on mapping columns)
      - other merge fields like «IdNo»
      - also a marker {Diagram} where the image should be inserted.
    """
    # Load mapping from external Excel
    if not os.path.exists(LOAD_MAPPING_FILE):
        raise FileNotFoundError(f"Load mapping file not found: {LOAD_MAPPING_FILE}")
    load_map = load_load_mapping(LOAD_MAPPING_FILE)
    
    # Read data Excel
    df = pd.read_excel(excel_path, sheet_name=sheet_name)
    df = df.fillna('').astype(str)
    records = df.to_dict(orient='records')
    
    if not records:
        raise ValueError("No data rows found in shackle sheet.")
    
    temp_dir = tempfile.mkdtemp()
    output_path = os.path.join(temp_dir, output_filename)
    
    final_doc = Document()
    
    for idx, row in enumerate(records):
        # Determine diagram based on Equipment and Type
        equipment = row.get('Equipment', 'bow shackle')
        type_ = row.get('Type', 'bolt and nut' or 'bolt & nut')
        diagram_path = get_diagram_path(equipment, type_)
        if not os.path.exists(diagram_path):
            print(f"Warning: Diagram not found at {diagram_path}, using fallback.")
            diagram_path = "shackles/Bow-BnN.png"
        
        # Compute pinCheck value
        pin_check = "N/A" if type_.lower() == "screw pin" else "ok"
        row['pinCheck'] = pin_check   # add to row for replacement
        
        # Get load value (assume column name 'Load_Tonnes')
        try:
            load = float(row.get('WLL', 0))
        except:
            load = 0.0
        
        # Look up mapping values (use closest or default)
        if load in load_map:
            mapping_vals = load_map[load]
        else:
            # Find nearest load? For simplicity, use first mapping or empty
            mapping_vals = load_map.get(list(load_map.keys())[0], {})
        
        # Clone template
        doc = Document(template_path)
        
        # ---- Replace all placeholders (both from row and from mapping) ----
        # First replace mapping placeholders (like «SWL», «Diameter»)
        for placeholder, value in mapping_vals.items():
            # The placeholder name in Word should match the Excel column name from mapping file
            # e.g., column 'SWL' becomes «SWL»
            replace_in_doc(doc, f"«{placeholder}»", str(value))
        
        # Then replace row data (including pinCheck and regular fields)
        for field_name, value in row.items():
            replace_in_doc(doc, f"«{field_name}»", str(value))
        
        # ---- Insert diagram at {Diagram} marker ----
        for paragraph in doc.paragraphs:
            if "{Diagram}" in paragraph.text:
                paragraph.clear()
                run = paragraph.add_run()
                run.add_picture(diagram_path)  # adjust size
                break
        
        # ---- Append to final document ----
        if idx > 0:
            final_doc.add_page_break()
        for element in doc.element.body:
            final_doc.element.body.append(element)
    
    final_doc.save(output_path)
    return output_path

def replace_in_doc(doc, placeholder, value):
    """Replace all occurrences of placeholder in paragraphs and tables."""
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, value)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if placeholder in cell.text:
                    cell.text = cell.text.replace(placeholder, value)